import io
import json
import re
import base64
import logging
import posixpath
import zipfile
import pandas as pd
from lxml import etree
from dotenv import load_dotenv

from typing import Optional
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware

load_dotenv()

from langchain_experimental.text_splitter import SemanticChunker
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain_core.messages import HumanMessage, SystemMessage
from langchain_core.documents import Document as LCDocument
from langchain_community.vectorstores import FAISS

from pypdf import PdfReader
from docx import Document as DocxDocument
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.shared import Pt
from pathlib import Path
from docx.oxml.ns import qn

BASE_DIR = Path(__file__).resolve().parent
AOR_TEMPLATE_PATH = BASE_DIR / "templates" / "Sample AOR Template.docx"

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.get("/")
async def root():
    return {"status": "ok", "version": "test-commit-1"}


# =========================================================
# EXTRACTION FUNCTIONS
# =========================================================

def extract_txt(contents: bytes, filename: str):
    text = contents.decode("utf-8", errors="ignore")
    return [
        LCDocument(
            page_content=text,
            metadata={
                "filename": filename,
                "source_type": "txt",
                "content_type": "plain_text"
            }
        )
    ]


def extract_pdf(contents: bytes, filename: str):
    docs = []
    try:
        pdf_stream = io.BytesIO(contents)
        pdf_reader = PdfReader(pdf_stream)
        for page_number, page in enumerate(pdf_reader.pages, start=1):
            page_text = page.extract_text() or ""
            if page_text.strip():
                docs.append(
                    LCDocument(
                        page_content=page_text,
                        metadata={
                            "filename": filename,
                            "source_type": "pdf",
                            "page": page_number,
                            "content_type": "page_text"
                        }
                    )
                )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to parse PDF: {str(e)}")
    return docs

RELS_NS = "http://schemas.openxmlformats.org/package/2006/relationships"


def _strip_dangling_relationships(contents: bytes) -> bytes:
    """Remove .rels entries pointing at parts missing from the archive.

    Some tools re-save .docx files without cleaning up relationship
    references (e.g. to word/people.xml for @mentions), which makes
    python-docx raise a KeyError when it tries to walk every part.
    """
    zin = zipfile.ZipFile(io.BytesIO(contents))
    names = set(zin.namelist())
    buffer = io.BytesIO()

    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)

            if item.filename.endswith(".rels"):
                base_dir = posixpath.dirname(posixpath.dirname(item.filename))
                tree = etree.fromstring(data)

                for rel in tree.findall(f"{{{RELS_NS}}}Relationship"):
                    if rel.get("TargetMode") == "External":
                        continue
                    target = posixpath.normpath(
                        posixpath.join(base_dir, rel.get("Target"))
                    )
                    if target not in names:
                        tree.remove(rel)

                data = etree.tostring(
                    tree, xml_declaration=True, encoding="UTF-8", standalone=True
                )

            zout.writestr(item, data)

    return buffer.getvalue()

def strip_cost_breakdown_narrative(text: str) -> str:
    """Drop CAPEX/OPEX narrative lines that duplicate the cost breakdown
    table (section labels, itemised line items, Sub Total/Contingency/
    Grand Total figures). The table is populated separately from
    capex_items/opex_items, so repeating them as text — whether in the
    docx "Estimated Costs" paragraph or the chat/preview response — is
    redundant, and the LLM's own arithmetic for these figures (e.g. Sub
    Total + 5% contingency) is not reliably correct.
    """
    label_line = re.compile(r"^\*{0,2}(capex|opex)\*{0,2}\s*:?\*{0,2}$", re.IGNORECASE)
    item_line = re.compile(r"^\d+\.\s.*\$[\d,]+(\.\d+)?\s*$")
    total_line = re.compile(
        r"^(sub\s*total|contingency\b.*|grand\s*total)\s*:?.*\$[\d,]+(\.\d+)?\s*$",
        re.IGNORECASE,
    )
    # Some drafts render the cost breakdown as a markdown table (e.g.
    # "| Description | Estimated Cost ($) |") instead of the plain-text
    # forms above — that duplicates the real docx table just as much.
    markdown_table_line = re.compile(r"^\|.*\|\s*$")
    markdown_table_separator = re.compile(r"^\|?[\s:|-]+\|[\s:|-]*\|?$")
    # The LLM's draft sometimes echoes back the template's own bracketed
    # authoring instruction (e.g. "[Insert cost breakdown table here]")
    # instead of omitting it — strip that substring out of the line
    # (keeping any real sentence text around it) rather than requiring
    # the whole line to be nothing but the placeholder.
    placeholder_bracket = re.compile(r"\s*\[insert[^\]]*\]", re.IGNORECASE)

    kept_lines = []
    for raw_line in text.splitlines():
        had_content = bool(raw_line.strip())
        line = placeholder_bracket.sub("", raw_line).rstrip()
        stripped = line.strip()
        if had_content and not stripped:
            # The whole line was just the placeholder — drop it entirely
            # rather than keeping an empty line in its place.
            continue
        if (
            label_line.match(stripped)
            or item_line.match(stripped)
            or total_line.match(stripped)
            or markdown_table_line.match(stripped)
            or markdown_table_separator.match(stripped)
        ):
            continue
        kept_lines.append(line)
    return "\n".join(kept_lines)

def compute_cost_totals(inputs: dict, metrics: dict) -> dict:
    """Single source of truth for CAPEX/OPEX totals, shared by the docx cost
    table and the chat/preview narrative so the two never disagree.
    """
    capex_items = inputs.get("capex_items") or []
    opex_items = inputs.get("opex_items") or []

    def sum_items(items):
        amounts = [
            item.get("amount") for item in items
            if isinstance(item.get("amount"), (int, float))
        ]
        return sum(amounts) if amounts else None

    # Total CAPEX/OPEX must equal the sum of the itemised rows — using the
    # separately-extracted top-level capex/opex figure instead (which may
    # already include contingency, e.g. a source Grand Total) would make the
    # total not add up to its own line items. Only fall back to that figure
    # when no items were itemised.
    capex_from_items = sum_items(capex_items)
    opex_from_items = sum_items(opex_items)
    capex_total = capex_from_items if capex_from_items is not None else inputs.get("capex")
    opex_total = opex_from_items if opex_from_items is not None else inputs.get("opex")

    total_cost = (
        capex_total + opex_total
        if isinstance(capex_total, (int, float)) and isinstance(opex_total, (int, float))
        else (metrics or {}).get("total_cost")
    )

    # Grand Total is always computed as Sub Total + 5% contingency, never
    # extracted from the document — extraction proved unreliable (it could
    # pick up a rounded administrative "Say" figure, or an unrelated number
    # entirely), whereas the 5% rate is fixed by the template's own heading
    # ("Grand Total (With 5% contingency)"), so computing it is both simpler
    # and guaranteed correct. "Grand Total" always means the combined
    # CAPEX+OPEX total with contingency, wherever it appears in the
    # document — there is no separate OPEX-only grand total.
    grand_total = total_cost * 1.05 if isinstance(total_cost, (int, float)) else None

    return {
        "capex_items": capex_items,
        "opex_items": opex_items,
        "capex_total": capex_total,
        "opex_total": opex_total,
        "total_cost": total_cost,
        "grand_total": grand_total,
    }

def build_cost_breakdown_bullets(inputs: dict, metrics: dict) -> str:
    """Deterministic bullet-point cost breakdown for the chat/preview
    response, built from the same totals as the docx cost table (via
    compute_cost_totals) so the two always agree — the LLM only ever
    writes the short intro sentence, never the figures themselves.
    """
    totals = compute_cost_totals(inputs, metrics or {})

    def money(val):
        if val is None:
            return "Not available"
        try:
            return f"${float(val):,.0f}"
        except (ValueError, TypeError):
            return str(val)

    lines = []
    if totals["capex_items"]:
        lines.append("CAPEX")
        for item in totals["capex_items"]:
            lines.append(f"- {item.get('description', '')}: {money(item.get('amount'))}")
    lines.append(f"- Total CAPEX: {money(totals['capex_total'])}")

    if totals["opex_items"]:
        lines.append("")
        lines.append("OPEX")
        for item in totals["opex_items"]:
            lines.append(f"- {item.get('description', '')}: {money(item.get('amount'))}")
    lines.append(f"- Total OPEX: {money(totals['opex_total'])}")

    lines.append(f"- Total Cost (CAPEX + OPEX): {money(totals['total_cost'])}")
    if totals["grand_total"] is not None:
        lines.append(f"- Grand Total (With 5% contingency): {money(totals['grand_total'])}")
    return "\n".join(lines)

def build_nev_bullets(metrics: dict) -> str:
    """Deterministic bullet-point NEV breakdown for the chat/preview
    response, built from the same computed_metrics as the docx NEV table —
    the LLM only ever writes the short qualitative sentence, never the
    figures themselves.
    """
    metrics = metrics or {}

    def money(val):
        if val is None:
            return "Not available – required information not found"
        try:
            return f"${float(val):,.2f}"
        except (ValueError, TypeError):
            return str(val)

    def ratio(val, suffix: str = ""):
        if val is None:
            return "Not available – required information not found"
        try:
            return f"{float(val):,.4f}{suffix}"
        except (ValueError, TypeError):
            return str(val)

    return "\n".join([
        f"- Present Value of Benefits: {money(metrics.get('total_benefits_pv'))}",
        f"- Present Value of Costs: {money(metrics.get('total_costs_pv'))}",
        f"- Net Present Value: {money(metrics.get('net_present_value'))}",
        f"- Benefit-Cost Ratio: {ratio(metrics.get('benefit_cost_ratio'))}",
        f"- Annual Manpower Impact: {ratio(metrics.get('annual_manpower_impact_fte'), ' FTE')}",
        f"- Annual Benefit: {money(metrics.get('annual_benefit'))}",
    ])

def insert_deterministic_breakdown(text: str, heading_names: list, breakdown_text: str) -> str:
    """Insert deterministic bullet text right after a section's intro
    sentence (the first non-empty line following one of heading_names), so
    the chat/preview response shows the same figures as the docx tables
    instead of relying on the LLM to restate them.
    """
    if not breakdown_text:
        return text

    targets = {
        re.sub(r"\s+", " ", name.strip().lower()).rstrip(":")
        for name in heading_names
    }

    lines = text.splitlines()
    result_lines = []
    i = 0
    inserted = False
    while i < len(lines):
        line = lines[i]
        result_lines.append(line)
        if not inserted:
            candidate = re.sub(r"^#{1,6}\s*", "", line.strip())
            candidate = re.sub(r"\s+", " ", candidate.replace("*", "")).rstrip(":").lower()
            if candidate in targets:
                i += 1
                while i < len(lines) and not lines[i].strip():
                    result_lines.append(lines[i])
                    i += 1
                if i < len(lines):
                    result_lines.append(lines[i])
                    i += 1
                result_lines.append("")
                result_lines.extend(breakdown_text.splitlines())
                inserted = True
                continue
        i += 1

    return "\n".join(result_lines)

def strip_nev_metrics_narrative(text: str) -> str:
    """Drop the individual NEV metric lines (Present Value of Benefits/Costs,
    Net Present Value, Benefit-Cost Ratio, Annual Manpower Impact, Annual
    Benefit) that duplicate the deterministic NEV table. Keeps only the
    LLM's qualitative assessment sentence(s).
    """
    metric_line = re.compile(
        r"^\*{0,2}-?\s*(present value of benefits|present value of costs|net present value|"
        r"benefit-cost ratio|annual manpower impact|annual benefit)\s*:.*$",
        re.IGNORECASE,
    )
    kept_lines = [
        line for line in text.splitlines()
        if not metric_line.match(line.strip())
    ]
    return "\n".join(kept_lines)

def build_result_docx(result_text: str, inputs: dict = None, metrics: dict = None) -> bytes:
    """
    Populate the existing AOR template sections using the drafted AOR text.
    The template's layout, tables, headers, footers and annexes are retained.
    """

    doc = DocxDocument(str(AOR_TEMPLATE_PATH))

    SECTION_ALIASES = {
        "purpose": "purpose",
        "background": "background",
        "need for deployment": "need",
        "need for deployment of the sample system": "need",
        "scope of work": "scope",
        "estimated costs": "costs",
        "proposed budget": "costs",
        "net economic value (nev) analysis and manpower capitalisation": "nev",
        "net economic value analysis and manpower capitalisation": "nev",
        "funding": "funding",
        "availability of funds": "funding",
        "approving authority": "authority",
        "approval": "approval",
    }

    def normalise_heading(text: str) -> str:
        text = text.strip()
        # The drafting LLM sometimes formats headings as markdown (e.g.
        # "## Purpose") even though the template's own headings are plain
        # text — strip that prefix first so heading detection isn't
        # silently defeated by markdown formatting.
        text = re.sub(r"^#{1,6}\s*", "", text)
        text = re.sub(r"^\d+[\.\s]+", "", text)
        text = text.replace("**", "")
        text = text.rstrip(":")
        return re.sub(r"\s+", " ", text).strip().lower()

    # All headings that can appear in the template — used to recognise where
    # one section ends and the next begins when replacing a section's body.
    # SECTION_ALIASES alone isn't enough: it models how the *drafting LLM*
    # phrases headings (e.g. "Need for Deployment"), but the template's own
    # placeholder headings (e.g. "Need for <xx>") don't always match that
    # phrasing, so scanning would otherwise run past them, swallow the next
    # heading as if it were body content, and delete it.
    ALL_HEADING_NAMES = {
        normalise_heading(name) for name in [
            "Purpose",
            "Background",
            "Need for <xx>",
            "Need For Deployment",
            "Need for Deployment of the Sample System",
            "Scope of Work",
            "Proposed Budget",
            "Estimated Costs",
            "Net Economic Value (NEV) Analysis and Manpower Capitalisation",
            "Net Economic Value Analysis and Manpower Capitalisation",
            "Availability of Funds",
            "Funding",
            "Approving Authority",
            "Approval",
        ]
    } | set(SECTION_ALIASES.keys())

    def fmt(val):
        if val == "" or val is None:
            return ""
        try:
            return f"${float(val):,.0f}"
        except (ValueError, TypeError):
            return str(val)

    def populate_budget_table(inputs: dict, metrics: dict):
        print("DEBUG inputs:", inputs)
        print("DEBUG metrics:", metrics)

        totals = compute_cost_totals(inputs, metrics)
        section_items = {
            "capex": totals["capex_items"],
            "opex": totals["opex_items"],
        }
        capex_total = totals["capex_total"]
        opex_total = totals["opex_total"]
        total_cost = totals["total_cost"]
        grand_total = totals["grand_total"]

        print("DEBUG values:", capex_total, opex_total, grand_total, total_cost)

        def set_cell_text(cell, text):
            for para in cell.paragraphs:
                for run in para.runs:
                    run.text = ""
            if cell.paragraphs[0].runs:
                run = cell.paragraphs[0].runs[0]
                run.text = text
            else:
                run = cell.paragraphs[0].add_run(text)
            run.font.name = "Arial"
            run.font.size = Pt(12)

        desc_col = None
        cost_col = None

        for i, table in enumerate(doc.tables):
            header_cells = [cell.text.strip().lower() for cell in table.rows[0].cells]
            header_text = " ".join(header_cells)
            print(f"DEBUG table {i} header: {header_text}")

            table_desc_col = next(
                (idx for idx, h in enumerate(header_cells) if "description" in h),
                None
            )
            table_cost_col = next(
                (idx for idx, h in enumerate(header_cells) if "amount" in h),
                None
            )
            if table_cost_col is None:
                table_cost_col = next(
                    (idx for idx, h in enumerate(header_cells) if "cost" in h and "unit" not in h),
                    None
                )

            # A table may continue a previous table's columns without repeating
            # its own header row (e.g. split across a page break) — in that
            # case keep using the last table's column positions.
            if table_desc_col is not None or table_cost_col is not None:
                desc_col, cost_col = table_desc_col, table_cost_col

            current_section = None
            item_index = 0

            for row in table.rows:
                first_cell = row.cells[0].text.strip().lower().rstrip(":")
                last_cell_index = len(row.cells) - 1
                # Prefer the header-derived cost column: some of the template's
                # tables have broken row-spanning merges around their total
                # rows, where the last cell in row.cells resolves to the wrong
                # (or a shared/aliased) cell. The header column is stable.
                total_col = cost_col if cost_col is not None else last_cell_index

                print(f"DEBUG row first_cell: '{first_cell}', last_cell_index: {last_cell_index}")

                if "total capex" in first_cell:
                    print("DEBUG matched total capex, writing:", fmt(capex_total))
                    set_cell_text(row.cells[total_col], fmt(capex_total))
                elif "total opex" in first_cell:
                    print("DEBUG matched total opex, writing:", fmt(opex_total))
                    set_cell_text(row.cells[total_col], fmt(opex_total))
                elif "total cost" in first_cell:
                    print("DEBUG matched total cost, writing:", fmt(total_cost))
                    set_cell_text(row.cells[total_col], fmt(total_cost))
                elif "grand total" in first_cell:
                    print("DEBUG matched grand total, writing:", fmt(grand_total))
                    set_cell_text(row.cells[total_col], fmt(grand_total))
                elif "capex" in first_cell:
                    current_section = "capex"
                    item_index = 0
                elif "opex" in first_cell:
                    current_section = "opex"
                    item_index = 0
                elif (
                    current_section is not None
                    and desc_col is not None
                    and cost_col is not None
                    and first_cell.isdigit()
                ):
                    items = section_items[current_section]
                    if item_index < len(items):
                        item = items[item_index]
                        print(f"DEBUG matched {current_section} item {item_index}, writing:", item)
                        set_cell_text(row.cells[desc_col], str(item.get("description") or ""))
                        set_cell_text(row.cells[cost_col], fmt(item.get("amount", "")))
                    item_index += 1

        return total_cost

    def parse_aor_sections(text: str):
        sections = {}
        title_lines = []
        current_section = None
        current_lines = []

        def save_current_section():
            if current_section and current_lines:
                sections[current_section] = "\n".join(
                    line for line in current_lines if line.strip()
                ).strip()

        for raw_line in text.splitlines():
            line = raw_line.strip()

            if not line:
                continue

            normalised = normalise_heading(line)
            detected_section = SECTION_ALIASES.get(normalised)

            if detected_section:
                save_current_section()
                current_section = detected_section
                current_lines = []
                continue

            if current_section is None:
                title_lines.append(line)
            else:
                current_lines.append(line)

        save_current_section()

        title = title_lines[0] if title_lines else ""
        title = re.sub(r"^#{1,6}\s*", "", title.strip())
        title = title.replace("**", "").strip()

        return title, sections

    # Debug/test code — must be at module level, not inside the function
    title, sections = parse_aor_sections(result_text)
    print("PARSED SECTIONS:", list(sections.keys()))
    print("NEED SECTION:", sections.get("need", "MISSING"))
    print("NEV SECTION:", sections.get("nev", "MISSING"))

    def replace_paragraph_text(paragraph, new_text: str):
        """Replace paragraph contents while retaining paragraph formatting."""
        p_elem = paragraph._p
        for child in list(p_elem):
            if child.tag == qn("w:pPr"):
                continue
            p_elem.remove(child)

        lines = new_text.splitlines()
        for index, line in enumerate(lines):
            if index > 0:
                break_run = paragraph.add_run()
                break_run.font.name = "Arial"
                break_run.font.size = Pt(12)
                break_run.add_break()
            cleaned_line = re.sub(r"^[-*]\s+", "", line.strip())
            run = paragraph.add_run(cleaned_line)
            run.font.name = "Arial"
            run.font.size = Pt(12)

    def find_heading_index(heading_names):
        names = {
            normalise_heading(name)
            for name in heading_names
        }

        for index, paragraph in enumerate(doc.paragraphs):
            paragraph_heading = normalise_heading(paragraph.text)

            if paragraph_heading in names:
                return index

        return None

    def replace_section_body(heading_names, new_text: str, stop_before_table: bool = False):
        if not new_text:
            return

        heading_index = find_heading_index(heading_names)
        if heading_index is None:
            return

        paragraphs = doc.paragraphs
        first_content_index = None
        paragraphs_to_remove = []

        for index in range(heading_index + 1, len(paragraphs)):
            paragraph = paragraphs[index]
            paragraph_text = paragraph.text.strip()
            normalised = normalise_heading(paragraph_text)

            if paragraph_text and normalised in ALL_HEADING_NAMES:
                break

            # If stop_before_table is set, the paragraph immediately before
            # the table ends the scan — but it's still just another content
            # paragraph for replace/remove purposes below (e.g. a template
            # placeholder like "[Insert cost breakdown table here]" sitting
            # between the intro sentence and the table must still be queued
            # for removal, not left in place just because it's adjacent to
            # the table).
            is_before_table = False
            if stop_before_table:
                p_elem = paragraph._p
                next_sibling = p_elem.getnext()
                is_before_table = next_sibling is not None and next_sibling.tag == qn("w:tbl")

            if paragraph_text and first_content_index is None:
                first_content_index = index
                replace_paragraph_text(paragraph, new_text)
            elif paragraph_text:
                paragraphs_to_remove.append(paragraph)

            if is_before_table:
                break

        for paragraph in paragraphs_to_remove:
            p_elem = paragraph._p
            p_elem.getparent().remove(p_elem)

    # Outside replace_section_body entirely
    title, sections = parse_aor_sections(result_text)

    # Replace the template title.
    if title:
        for paragraph in doc.paragraphs:
            paragraph_text = paragraph.text.strip()

            if "<TITLE>" in paragraph_text.upper():
                replace_paragraph_text(
                    paragraph,
                    title.upper()
                )
                break

    # Populate each corresponding template section.
    replace_section_body(
        ["Purpose"],
        sections.get("purpose", "")
    )

    replace_section_body(
        ["Background"],
        sections.get("background", "")
    )

    replace_section_body(
        [
            "Need for <xx>",
            "Need For Deployment",
            "Need for Deployment of the Sample System"
        ],
        sections.get("need", "")
    )

    # Fill the "<xx>" placeholder in the "Need for <xx>" heading itself
    # (not just its body, replaced above) with the project title, so it
    # reads e.g. "Need for Sample Digital Booking System" instead of
    # leaving the raw template placeholder in the output.
    if inputs and inputs.get("project_title"):
        for paragraph in doc.paragraphs:
            if normalise_heading(paragraph.text) == normalise_heading("Need for <xx>"):
                for run in paragraph.runs:
                    if "<xx>" in run.text:
                        run.text = run.text.replace("<xx>", inputs["project_title"])
                break

    replace_section_body(
        ["Scope of Work"],
        sections.get("scope", "")
    )

    replace_section_body(
        ["Proposed Budget", "Estimated Costs"],
        strip_cost_breakdown_narrative(sections.get("costs", "")),
        stop_before_table=True
    )

    # The template has no "Net Economic Value" heading of its own to reuse
    # (unlike the other sections above), so it's built from scratch here —
    # a heading, the LLM's short qualitative assessment sentence, and a
    # table of the computed metrics — and inserted right before the
    # "Availability of Funds" heading (the section that logically follows
    # the budget in document order).
    if metrics:
        nev_anchor_index = find_heading_index(["Availability of Funds", "Funding"])
        if nev_anchor_index is not None:
            nev_anchor = doc.paragraphs[nev_anchor_index]

            def insert_before(text: str, bold: bool = False, size: int = 12):
                new_paragraph = nev_anchor.insert_paragraph_before(text)
                for run in new_paragraph.runs:
                    run.bold = bold
                    run.font.name = "Arial"
                    run.font.size = Pt(size)
                return new_paragraph

            insert_before(
                "Net Economic Value (NEV) Analysis and Manpower Capitalisation",
                bold=True,
                size=14
            )
            insert_before("")

            nev_intro = strip_nev_metrics_narrative(sections.get("nev", "")).strip()
            if nev_intro:
                insert_before(nev_intro)
                insert_before("")

            def fmt_ratio(val, suffix: str = ""):
                if val is None:
                    return "Not available – required information not found"
                try:
                    return f"{float(val):,.4f}{suffix}"
                except (ValueError, TypeError):
                    return str(val)

            def fmt_or_na(val):
                # 2dp (not the 0dp used for CAPEX/OPEX) so this matches the
                # precision of the deterministic NEV bullets appended to the
                # chat/preview response — present values legitimately carry
                # meaningful cents.
                if val is None:
                    return "Not available – required information not found"
                try:
                    return f"${float(val):,.2f}"
                except (ValueError, TypeError):
                    return str(val)

            nev_rows = [
                ("Present Value of Benefits", fmt_or_na(metrics.get("total_benefits_pv"))),
                ("Present Value of Costs", fmt_or_na(metrics.get("total_costs_pv"))),
                ("Net Present Value", fmt_or_na(metrics.get("net_present_value"))),
                ("Benefit-Cost Ratio", fmt_ratio(metrics.get("benefit_cost_ratio"))),
                ("Annual Manpower Impact (FTE)", fmt_ratio(metrics.get("annual_manpower_impact_fte"), " FTE")),
                ("Annual Benefit", fmt_or_na(metrics.get("annual_benefit"))),
            ]

            nev_table = doc.add_table(rows=len(nev_rows) + 1, cols=2)
            nev_table.style = "Table Grid"
            nev_table.rows[0].cells[0].text = "Metric"
            nev_table.rows[0].cells[1].text = "Value"
            for row, (label, value) in zip(nev_table.rows[1:], nev_rows):
                row.cells[0].text = label
                row.cells[1].text = value
            for row in nev_table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.name = "Arial"
                            run.font.size = Pt(12)

            nev_anchor._p.addprevious(nev_table._tbl)
            insert_before("")

    replace_section_body(
        ["Availability of Funds", "Funding"],
        sections.get("funding", "")
    )

    # The Approving Authority sentence is left as the template's own
    # formal wording ("In accordance with Paragraph 8.4.3, ... <xx> is the
    # appropriate approving authority for operating expenditure up to
    # <xx>.") — its two "<xx>" placeholders are filled in below, rather
    # than replacing the whole sentence with the LLM's free-form phrasing.

    replace_section_body(
        ["Approval"],
        sections.get("approval", "")
    )

    # Populate budget table if structured data is available
    if inputs and metrics:
        table_total_cost = populate_budget_table(inputs, metrics)

        # Fill the "<xx>" placeholder in the "proposed budget of <xx>
        # (without 5% contingency)" sentence with the same total cost
        # (CAPEX + OPEX, before contingency) shown in the table, so the
        # narrative and the table never disagree.
        for paragraph in doc.paragraphs:
            if "proposed budget of" in paragraph.text.lower() and "without 5% contingency" in paragraph.text.lower():
                for run in paragraph.runs:
                    if run.text.strip() == "<xx>":
                        run.text = run.text.replace("<xx>", fmt(table_total_cost))
                    # This sentence is left as the template's own wording
                    # (not rebuilt via replace_paragraph_text), so its font
                    # is normalised here to match the rest of the body text.
                    run.font.name = "Arial"
                    run.font.size = Pt(12)

    # Fill the two "<xx>" placeholders in the Approving Authority sentence
    # with the authority name and spending threshold for the grand total.
    # The sentence already ends "... up to <xx>.", so the threshold's own
    # "Up to " prefix is dropped for the second placeholder to avoid
    # "up to Up to S$5 million".
    if inputs:
        approval_info = determine_approving_authority(
            compute_cost_totals(inputs, metrics or {})["grand_total"]
        )
        if approval_info:
            authority_name, threshold_label = approval_info
            threshold_amount = re.sub(r"^(up to|above)\s+", "", threshold_label, flags=re.IGNORECASE)
            xx_values = [authority_name, threshold_amount]
            for paragraph in doc.paragraphs:
                if "appropriate approving authority" in paragraph.text.lower():
                    xx_index = 0
                    for run in paragraph.runs:
                        if run.text.strip() == "<xx>" and xx_index < len(xx_values):
                            run.text = run.text.replace("<xx>", xx_values[xx_index])
                            xx_index += 1
                        # This sentence is left as the template's own
                        # wording (not rebuilt via replace_paragraph_text),
                        # so its font is normalised here to match the rest
                        # of the body text.
                        run.font.name = "Arial"
                        run.font.size = Pt(12)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

def _load_docx(contents: bytes) -> DocxDocument:
    try:
        return DocxDocument(io.BytesIO(contents))
    except KeyError:
        return DocxDocument(io.BytesIO(_strip_dangling_relationships(contents)))


def extract_docx(contents: bytes, filename: str):
    docs = []
    try:
        doc = _load_docx(contents)

        paragraphs = [
            para.text.strip()
            for para in doc.paragraphs
            if para.text.strip()
        ]
        if paragraphs:
            docs.append(
                LCDocument(
                    page_content="\n".join(paragraphs),
                    metadata={
                        "filename": filename,
                        "source_type": "docx",
                        "content_type": "paragraphs"
                    }
                )
            )
        for table_index, table in enumerate(doc.tables, start=1):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    rows.append(" | ".join(cells))
            if rows:
                docs.append(
                    LCDocument(
                        page_content="\n".join(rows),
                        metadata={
                            "filename": filename,
                            "source_type": "docx",
                            "content_type": "table",
                            "table_index": table_index
                        }
                    )
                )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to parse Word Document: {str(e)}")
    return docs


def _cost_table_snippet(table, max_rows: int = 3, max_len: int = 150) -> str:
    """Return a short snippet of a table's content if it looks like it
    discusses cost/budget figures (CAPEX/OPEX labels or dollar amounts),
    otherwise an empty string.

    Requires more than one row: a single-row table is itself just another
    heading label (e.g. "12 | Estimated Costs", the SAME two-column
    heading-table pattern used throughout these documents) — not an
    actual cost breakdown — and would otherwise trigger on the word
    "Costs" in a neighboring section's own title, wrongly making it look
    like cost data follows.
    """
    if len(table.rows) < 2:
        return ""
    parts = []
    has_signal = False
    for row in table.rows[:max_rows]:
        cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
        if not cells:
            continue
        row_text = " ".join(cells)
        parts.append(row_text)
        if re.search(r"capex|opex|cost|budget|\$|\d{3,}", row_text, re.IGNORECASE):
            has_signal = True
    return " | ".join(parts)[:max_len] if has_signal else ""


def extract_candidate_headings(contents: bytes, include_tables: bool = True, with_context: bool = False) -> list:
    """Pull short, standalone paragraphs and table rows that likely function
    as section headings. Many AOR documents lay out numbered section titles
    as a two-column table row (e.g. "10 | Scope of Work") rather than a
    plain paragraph, and bold a heading rather than apply Word's built-in
    Heading style — so paragraph.style alone isn't reliable. Length is used
    as a rough heuristic instead, and the LLM is left to judge which
    candidates are actually headings.

    `include_tables` can be disabled for documents (like the reference AOR
    template) where headings are all plain paragraphs and scanning tables
    would only add noise (row numbers, merged-cell duplication, etc.) from
    cost-breakdown tables.

    `with_context`, when a table-row heading comes from a small (<=2 row)
    standalone heading table, peeks at the table immediately following it
    for cost/budget signal (e.g. "8 | Estimated Costs" followed by a
    separate CAPEX/OPEX breakdown table) and appends a short snippet. This
    lets the LLM recognize a differently-worded heading (e.g. "Estimated
    Costs" vs. a reference "Proposed Budget") as the same section when the
    content it precedes makes that unambiguous, even with no shared words.
    """
    return [entry["text"] for entry in build_heading_index(contents, include_tables, with_context)]


def build_heading_index(contents: bytes, include_tables: bool = True, with_context: bool = False) -> list:
    """Like extract_candidate_headings, but returns each candidate's
    document position alongside its text:
    {"id": str, "text": str, "paragraph": Paragraph} or
    {"id": str, "text": str, "row": <table row>}.

    IDs are positional/deterministic (P{i} = i-th paragraph, T{ti}R{ri} =
    ri-th row of the ti-th table), so they can be recomputed identically
    from the same document bytes across separately-loaded Document
    instances — this lets a structure-review LLM call (analysing one
    loaded copy) and a later document-edit step (which loads its own
    fresh copy to modify and save) agree on what an ID refers to, without
    sharing Python objects across the two steps.

    Loads its own Document from `contents`. If you already have a loaded
    Document that you intend to modify and save, use
    _heading_index_from_doc(doc, ...) instead so the returned
    paragraph/row objects belong to the instance you'll actually save —
    otherwise inserting relative to them has no effect on your saved file.
    """
    return _heading_index_from_doc(_load_docx(contents), include_tables, with_context)


def _iter_body_items(doc):
    """Yield ("paragraph"|"table", object) for each top-level paragraph and
    table in the document body, in TRUE document order. python-docx's
    doc.paragraphs / doc.tables are each internally ordered, but the
    interleaving between the two collections is lost when read separately
    — this walks the underlying XML body directly to recover it, which
    matters for telling whether a heading appears before or after an
    Annex boundary.
    """
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield "paragraph", Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield "table", Table(child, doc)


def _heading_index_from_doc(doc, include_tables: bool = True, with_context: bool = False) -> list:
    body_items = list(_iter_body_items(doc))
    # IDs stay P{i}/T{ti}R{ri} using the same numbering as doc.paragraphs[i]
    # / doc.tables[ti] (each collection counted independently, in its own
    # order) — only the ORDER entries are appended to `index` changes here,
    # so existing ids/lookups by id are unaffected.
    index = []
    p_count = 0
    t_count = 0
    for pos, (kind, item) in enumerate(body_items):
        if kind == "paragraph":
            i = p_count
            p_count += 1
            text = item.text.strip()
            if text and len(text) <= 80:
                index.append({"id": f"P{i}", "text": text, "paragraph": item})
        else:
            ti = t_count
            t_count += 1
            if not include_tables:
                continue
            for ri, row in enumerate(item.rows):
                row_text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text and len(row_text) <= 80:
                    text = row_text
                    if with_context and len(item.rows) <= 2:
                        next_table = next(
                            (nxt for kind2, nxt in body_items[pos + 1:] if kind2 == "table"),
                            None
                        )
                        if next_table is not None:
                            snippet = _cost_table_snippet(next_table)
                            if snippet:
                                text = f"{row_text} (followed by a table mentioning: {snippet})"
                    index.append({"id": f"T{ti}R{ri}", "text": text, "row": row, "table": item})
    return index


# Section-title synonyms confirmed from real documents. Applied directly to
# candidate text (rather than left as a rule for the LLM to cross-reference)
# since testing showed the LLM doesn't reliably apply a stated synonym rule
# on its own, even when it's explicit — annotating the candidate itself
# makes the match definitional instead of inferential.
_KNOWN_SECTION_SYNONYMS = [
    (re.compile(r"\bfunding\b", re.IGNORECASE), "Availability of Funds"),
    (re.compile(r"\bavailability of funds\b", re.IGNORECASE), "Funding"),
    (re.compile(r"\bestimated costs?\b", re.IGNORECASE), "Proposed Budget"),
    (re.compile(r"\bproposed budget\b", re.IGNORECASE), "Estimated Costs"),
    (re.compile(r"\bneed for\b", re.IGNORECASE), "Need For Deployment"),
]


def _annotate_synonym_text(text: str) -> str:
    extras = [
        synonym
        for pattern, synonym in _KNOWN_SECTION_SYNONYMS
        if pattern.search(text) and synonym.lower() not in text.lower()
    ]
    return f"{text} (also referred to as: {', '.join(extras)})" if extras else text


def annotate_known_synonyms(candidates: list) -> list:
    """Append known canonical synonym names directly onto matching
    candidate headings, so the structure-review LLM sees the equivalence
    as a fact about the text rather than a rule it has to apply itself.
    """
    return [_annotate_synonym_text(candidate) for candidate in candidates]


def annotate_heading_index_synonyms(index: list) -> list:
    """Like annotate_known_synonyms, but for build_heading_index() entries
    — updates each entry's "text" in place (id/paragraph/row untouched).
    """
    for entry in index:
        entry["text"] = _annotate_synonym_text(entry["text"])
    return index


_TEMPLATE_PLACEHOLDER_RE = re.compile(r"<[^>]{1,40}>")
_TEMPLATE_BARE_PLACEHOLDER_RE = re.compile(r"^<[^>]{1,40}>$")


def _clean_template_heading_candidates(candidates: list) -> list:
    """Filter the AOR template's raw candidate headings down to genuine
    section titles, dropping cover-page metadata and placeholder filler.

    Tuned to this specific template's conventions (cover-page block, then a
    bare <TITLE> placeholder, then the body sections) — re-verify against
    the actual file if the reference template is ever swapped out.
    """
    # Drop the cover-page block (page count / forum / EMS reference / date)
    # that precedes the title placeholder in this template's layout.
    cut_index = None
    for i, candidate in enumerate(candidates):
        if _TEMPLATE_BARE_PLACEHOLDER_RE.match(candidate.strip()):
            cut_index = i
            break
    working = candidates[cut_index + 1:] if cut_index is not None else candidates

    cleaned = []
    seen = set()
    for raw in working:
        text = re.sub(r"\s+", " ", raw.strip())
        if not text or ":" in text:
            # Drops table captions ("Table 1: Cost Breakdown Table") and
            # the signature block ("Prepared by:", "Vetted by:", etc.).
            continue
        stripped = re.sub(r"\s+", " ", _TEMPLATE_PLACEHOLDER_RE.sub("", text)).strip()
        if sum(ch.isalpha() for ch in stripped) < 3:
            # Pure filler like "<to fill up>" or a bare "<xx>".
            continue
        # Use the placeholder-stripped text (e.g. "Need for" rather than
        # "Need for <xx>") so the reference name reads as natural language
        # for the LLM comparison, instead of raw template placeholder syntax.
        text = stripped
        key = stripped.lower()
        if key in seen:
            continue
        seen.add(key)
        cleaned.append(text)
    return cleaned


def extract_excel(contents: bytes, filename: str):
    docs = []
    try:
        excel = pd.read_excel(io.BytesIO(contents), sheet_name=None)
        for sheet_name, df in excel.items():
            if df.empty:
                continue
            df.columns = [str(col).strip() for col in df.columns]
            docs.append(
                LCDocument(
                    page_content=df.to_csv(index=False),
                    metadata={
                        "filename": filename,
                        "source_type": "excel",
                        "sheet": sheet_name,
                        "content_type": "sheet_table",
                        "rows": len(df),
                        "columns": list(df.columns)
                    }
                )
            )
            for idx, row in df.iterrows():
                row_dict = row.dropna().to_dict()
                if row_dict:
                    docs.append(
                        LCDocument(
                            page_content=json.dumps(row_dict, ensure_ascii=False),
                            metadata={
                                "filename": filename,
                                "source_type": "excel",
                                "sheet": sheet_name,
                                "content_type": "row",
                                "row_index": idx + 1
                            }
                        )
                    )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to parse Excel Document: {str(e)}")
    return docs


def extract_documents(contents: bytes, filename: str):
    filename_lower = filename.lower()
    if filename_lower.endswith(".txt"):
        return extract_txt(contents, filename)
    if filename_lower.endswith(".pdf"):
        return extract_pdf(contents, filename)
    if filename_lower.endswith(".docx"):
        return extract_docx(contents, filename)
    if filename_lower.endswith((".xlsx", ".xls")):
        return extract_excel(contents, filename)
    raise HTTPException(
        status_code=400,
        detail="Unsupported file format. Please upload .txt, .pdf, .docx, .xlsx, or .xls files."
    )


# =========================================================
# SPLITTING
# =========================================================

def split_documents(raw_docs, embeddings):
    splitter = SemanticChunker(embeddings)
    result = []
    for doc in raw_docs:
        # Tables (and Excel rows) must stay intact — splitting them can
        # separate a section header (e.g. "CAPEX"/"OPEX") from the rows/
        # totals underneath it, causing the LLM to mislabel which figure
        # belongs to which section.
        if doc.metadata.get("content_type") in ("row", "table"):
            result.append(doc)
            continue
        chunks = splitter.create_documents(
            [doc.page_content],
            metadatas=[doc.metadata]
        )
        for i, chunk in enumerate(chunks, start=1):
            chunk.metadata["chunk_index"] = i
            result.append(chunk)
    return result


# =========================================================
# JSON AND NUMBER HELPERS
# =========================================================

def parse_json_from_llm(text: str):
    cleaned = text.strip()
    cleaned = cleaned.replace("```json", "")
    cleaned = cleaned.replace("```", "")
    cleaned = cleaned.strip()
    try:
        return json.loads(cleaned)
    except Exception:
        pass
    match = re.search(r"\{.*\}", cleaned, re.DOTALL)
    if match:
        return json.loads(match.group(0))
    raise ValueError("No valid JSON object found in LLM response.")


def to_number(value):
    if value is None:
        return None
    if isinstance(value, (int, float)):
        return value
    if isinstance(value, str):
        cleaned = value.strip()
        if cleaned.lower() in ["not found", "not found in document", "na", "n/a", "none", "null", ""]:
            return None
        cleaned = cleaned.replace("S$", "").replace("$", "").replace(",", "").replace("%", "")
        match = re.search(r"-?\d+(\.\d+)?", cleaned)
        if match:
            number = float(match.group(0))
            return int(number) if number.is_integer() else number
    return None


def clean_extracted_inputs(inputs: dict):
    numeric_fields = [
        "capex",
        "opex",
        "project_duration_years",
        "annual_productivity_time_savings_hours",
        "annual_manpower_impact_fte",
        "annual_benefit",
        "num_staff",
        "savings_duration_months",
        "man_hour_rate"
    ]
    cleaned = dict(inputs)
    for field in numeric_fields:
        cleaned[field] = to_number(cleaned.get(field))

    for key in ("capex_items", "opex_items"):
        items = cleaned.get(key) or []
        cleaned[key] = [
            {**item, "amount": to_number(item.get("amount"))}
            for item in items if isinstance(item, dict)
        ]

    return cleaned


# =========================================================
# ICT AOR CALCULATIONS
# =========================================================

def compute_aor_metrics(inputs: dict):
    metrics = {}

    capex = inputs.get("capex")
    opex = inputs.get("opex")
    project_duration_years = inputs.get("project_duration_years")
    annual_productivity_time_savings_hours = inputs.get("annual_productivity_time_savings_hours")
    annual_manpower_impact_fte = inputs.get("annual_manpower_impact_fte")
    annual_benefit = inputs.get("annual_benefit")
    savings_duration_months = inputs.get("savings_duration_months")
    man_hour_rate = inputs.get("man_hour_rate")

    # --- Sanity log ---
    metrics["debug_project_duration_years"] = project_duration_years  # 👈 here
    metrics["debug_capex"] = capex
    metrics["debug_opex"] = opex

    hours_per_fte = 1819
    r = 0.06  # 4% nominal discount + 2% inflation

    # --- PV factor (annuity, computed from project duration) ---
    if project_duration_years is not None and project_duration_years > 0:
        pv_factor = (1 - (1 + r) ** -project_duration_years) / r
        metrics["pv_factor"] = round(pv_factor, 4)
    else:
        metrics["pv_factor"] = None

    # --- Annual OPEX (OPEX spread evenly over project duration) ---
    if opex is not None and project_duration_years is not None and project_duration_years > 0:
        metrics["annual_opex"] = round(opex / project_duration_years, 2)
    else:
        metrics["annual_opex"] = None

    # --- Annualised productivity hours ---
    if annual_productivity_time_savings_hours is not None:
        duration_months = savings_duration_months if savings_duration_months else 12
        metrics["annualised_productivity_hours"] = round(
            annual_productivity_time_savings_hours * (12 / duration_months), 2
        )
    else:
        metrics["annualised_productivity_hours"] = None

    # --- Annual manpower impact (FTE) ---
    if annual_manpower_impact_fte is not None:
        metrics["annual_manpower_impact_fte"] = annual_manpower_impact_fte
    elif metrics["annualised_productivity_hours"] is not None:
        metrics["annual_manpower_impact_fte"] = round(
            metrics["annualised_productivity_hours"] / hours_per_fte, 4
        )
    else:
        metrics["annual_manpower_impact_fte"] = None

    # --- Annual benefit ---
    if annual_benefit is not None:
        metrics["annual_benefit"] = annual_benefit
    elif metrics["annual_manpower_impact_fte"] is not None and man_hour_rate is not None:
        metrics["annual_benefit"] = round(
            metrics["annual_manpower_impact_fte"] * hours_per_fte * man_hour_rate, 2
        )
    else:
        metrics["annual_benefit"] = None

    # --- Total cost (simple sum, for reference only) ---
    if capex is not None and opex is not None:
        metrics["total_cost"] = round(capex + opex, 2)
    elif opex is not None:
        metrics["total_cost"] = round(opex, 2)
    else:
        metrics["total_cost"] = None

    # --- Total benefits PV (annual benefit discounted as annuity) ---
    if metrics["annual_benefit"] is not None and metrics["pv_factor"] is not None:
        metrics["total_benefits_pv"] = round(
            metrics["annual_benefit"] * metrics["pv_factor"], 2
        )
    else:
        metrics["total_benefits_pv"] = None

    # --- Total costs PV ---
    # CAPEX is already a present value (one-off upfront cost)
    # OPEX is treated as an annuity discounted over project duration
    if capex is not None and metrics["annual_opex"] is not None and metrics["pv_factor"] is not None:
        metrics["total_costs_pv"] = round(
            capex + (metrics["annual_opex"] * metrics["pv_factor"]), 2
        )
    else:
        metrics["total_costs_pv"] = None

    # --- Net Present Value ---
    if metrics["total_benefits_pv"] is not None and metrics["total_costs_pv"] is not None:
        metrics["net_present_value"] = round(
            metrics["total_benefits_pv"] - metrics["total_costs_pv"], 2
        )
    else:
        metrics["net_present_value"] = None

    # --- Benefit-Cost Ratio ---
    if (
        metrics["total_benefits_pv"] is not None
        and metrics["total_costs_pv"] is not None
        and metrics["total_costs_pv"] != 0
    ):
        metrics["benefit_cost_ratio"] = round(
            metrics["total_benefits_pv"] / metrics["total_costs_pv"], 4
        )
    else:
        metrics["benefit_cost_ratio"] = None

    return metrics

def identify_missing_fields(inputs: dict, metrics: dict):
    missing = []

    # These must come from the document — cannot be derived
    doc_required = [
        "project_title",
        "purpose",
        "problem_statement",
        "capex",
        "opex",
        "project_duration_years",
    ]
    for field in doc_required:
        if inputs.get(field) is None:
            missing.append(field)

    # These can be derived — only flag as missing if computation also failed
    derived_required = {
        "annual_manpower_impact_fte": metrics.get("annual_manpower_impact_fte"),
        "annual_benefit": metrics.get("annual_benefit"),
    }
    for field, computed_value in derived_required.items():
        if inputs.get(field) is None and computed_value is None:
            missing.append(field)

    return missing

def determine_approving_authority(amount):
    if amount is None:
        return None

    if amount <= 6000:
        return (
            "Deputy Director / Senior Assistant Director",
            "Up to S$6,000"
        )

    elif amount <= 100000:
        return (
            "Division Head",
            "Up to S$100,000"
        )

    elif amount <= 250000:
        return (
            "Senior Director",
            "Up to S$250,000"
        )

    elif amount <= 500000:
        return (
            "Assistant/Deputy Director-General",
            "Up to S$500,000"
        )

    elif amount <= 1000000:
        return (
            "Director-General",
            "Up to S$1 million"
        )

    elif amount <= 5000000:
        return (
            "Management Committee",
            "Up to S$5 million"
        )

    elif amount <= 10000000:
        return (
            "Chairman",
            "Up to S$10 million"
        )

    else:
        return (
            "Authority",
            "Above S$10 million"
        )
# =========================================================
# DOCUMENT AMENDMENT
# =========================================================

    # These must come from the document — cannot be derived
    doc_required = ["capex", "opex", "project_duration_years"]
    for field in doc_required:
        if inputs.get(field) is None:
            missing.append(field)

    # These can be derived — only flag as missing if computation also failed
    derived_required = {
        "annual_manpower_impact_fte": metrics.get("annual_manpower_impact_fte"),
        "annual_benefit": metrics.get("annual_benefit"),
    }
    for field, computed_value in derived_required.items():
        if inputs.get(field) is None and computed_value is None:
            missing.append(field)

    return missing


# =========================================================
# DOCUMENT AMENDMENT
# =========================================================

FIELD_LABELS = {
    "project_title": "Project Title",
    "purpose": "Purpose",
    "problem_statement": "Problem Statement",
    "capex": "Capital Expenditure (CAPEX)",
    "opex": "Operating Expenditure (OPEX)",
    "project_duration_years": "Project Duration (years)",
    "annual_manpower_impact_fte": "Annual Manpower Impact (FTE)",
    "annual_benefit": "Annual Benefit ($)",
}

FIELD_PLACEHOLDERS = {
    "project_title": "[Project Title — enter the name of the ICT project]",
    "purpose": "[Purpose — enter what the ICT project seeks to achieve]",
    "problem_statement": "[Problem Statement — enter the gap, pain point, or problem being addressed]",
    "capex": "[CAPEX — enter the total capital expenditure, e.g. S$500,000]",
    "opex": "[OPEX — enter the total operating expenditure over the project duration, e.g. S$300,000]",
    "project_duration_years": "[Project Duration — enter the number of years, e.g. 3]",
    "annual_manpower_impact_fte": "[Annual Manpower Impact (FTE) — enter the estimated FTE impact, e.g. 0.5]",
    "annual_benefit": "[Annual Benefit — enter the estimated annual dollar benefit, e.g. S$50,000]",
}


def build_missing_information_section(missing_fields: list) -> str:
    """Deterministically render the "Missing Information" section, rather
    than leaving its wording/formatting to the LLM each time. Only called
    when missing_fields is non-empty — when nothing is missing, the section
    is omitted entirely rather than stating that fact.
    """
    bullet_lines = "\n".join(
        f"- {FIELD_LABELS.get(field, field.replace('_', ' ').title())}"
        for field in missing_fields
    )
    return (
        "5. Missing Information / Follow-up Required\n\n"
        "You are missing the following required information. Please include this in your AOR:\n"
        f"{bullet_lines}"
    )


_COST_SIGNAL_RE = re.compile(r"\b(capex|opex|budget|costs?)\b", re.IGNORECASE)
_COST_LABEL_RE = re.compile(r"\b(capex|opex|budget|costs?)\s*[:=]", re.IGNORECASE)
_COST_TABLE_ANNOTATION_RE = re.compile(r"\(followed by a table mentioning:", re.IGNORECASE)


def _looks_like_cost_declaration(text: str) -> bool:
    """True if `text` looks like it's actually declaring a cost/budget
    figure or heading — not just prose that happens to mention the word
    while discussing something else (e.g. "Based on the illustrative
    CAPEX and OPEX, this project has a negative NEV." mentions both terms
    but is NEV narrative, not a cost declaration).

    Accepts: a "Label:"/"Label=" declaration (e.g. "CAPEX: S$500,000."),
    the cost-table annotation, or a short heading/table-cell-like
    candidate (<=6 words) that contains the term — genuine headings and
    table cells are short; narrative sentences citing the term in passing
    are not.
    """
    if _COST_TABLE_ANNOTATION_RE.search(text):
        return True
    if not _COST_SIGNAL_RE.search(text):
        return False
    if _COST_LABEL_RE.search(text):
        return True
    return len(text.split()) <= 6


def resolve_cost_section_matches(section_matches: dict, heading_index: list) -> dict:
    """Replace the LLM's answer for any standard section that is itself
    about cost/budget (name contains capex/opex/budget/cost) with a fully
    deterministic determination, ignoring what the LLM matched.

    Testing repeatedly showed the LLM doesn't reliably avoid matching such
    a section to unrelated content that merely cites cost figures in
    passing (e.g. a Net Present Value narrative referencing CAPEX/OPEX
    numbers), even with explicit prompt instructions not to — so presence
    is instead decided directly in code: present only if some candidate,
    before any Annex boundary, looks like an actual cost declaration or
    heading (see _looks_like_cost_declaration), not just prose mentioning
    the term; otherwise missing.
    """
    annex_start = _find_annex_start(heading_index)
    main_body_entries = (
        heading_index[:annex_start] if annex_start is not None else heading_index
    )
    match_id = next(
        (entry["id"] for entry in main_body_entries if _looks_like_cost_declaration(entry["text"])),
        None
    )

    resolved = dict(section_matches)
    for section_name in section_matches:
        if _COST_SIGNAL_RE.search(section_name):
            resolved[section_name] = match_id
    return resolved


def deduplicate_reused_matches(section_matches: dict, heading_index: list) -> dict:
    """If the LLM matched the same candidate heading to more than one
    standard section — which its own instructions tell it not to do, but
    testing shows it doesn't always follow, especially for lexically
    similar section names like "Approval" vs. "Approving Authority" —
    keep the match only for whichever section's name has the most word
    overlap with the candidate's own text, and null out the rest.
    """
    entry_by_id = {entry["id"]: entry for entry in heading_index}
    sections_by_match = {}
    for section, matched_id in section_matches.items():
        if matched_id:
            sections_by_match.setdefault(matched_id, []).append(section)

    resolved = dict(section_matches)
    for matched_id, sections in sections_by_match.items():
        if len(sections) <= 1:
            continue
        candidate_text = entry_by_id.get(matched_id, {}).get("text", "").lower()
        candidate_words = set(re.findall(r"[a-z]+", candidate_text))

        def overlap(section_name):
            section_words = set(re.findall(r"[a-z]+", section_name.lower()))
            return len(section_words & candidate_words)

        best_section = max(sections, key=overlap)
        for section in sections:
            if section != best_section:
                resolved[section] = None
    return resolved


def build_structure_review_section(missing_sections: list, order_note: str) -> str:
    """Deterministically render the "Document Structure Review" section
    from the LLM's structured missing_sections/order_note output, rather
    than trusting free-form narrative formatting each time.
    """
    if not missing_sections:
        text = "All standard AOR sections appear to be present."
    else:
        bullet_lines = "\n".join(f"- {section}" for section in missing_sections)
        text = (
            "You are missing the following sections. Please include this in your AOR:\n"
            f"{bullet_lines}"
        )
    if order_note:
        text += f"\n\n{order_note}"
    return "6. Document Structure Review\n\n" + text


def _insert_highlighted_paragraph_before(anchor_element, parent, text: str) -> Paragraph:
    """Insert a new highlighted paragraph immediately before `anchor_element`
    (a paragraph's `_p`, or a table's `_tbl`). Repeated calls against the
    same fixed anchor naturally preserve insertion order (each new
    paragraph lands directly before the anchor, pushing earlier
    insertions further back), so no chaining/tracking is needed here.
    """
    new_p = OxmlElement("w:p")
    anchor_element.addprevious(new_p)
    new_paragraph = Paragraph(new_p, parent)
    run = new_paragraph.add_run(text)
    run.bold = True
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return new_paragraph


def build_amended_docx(
    contents: bytes,
    missing_fields: list,
    missing_sections: list = None,
    section_matches: dict = None,
    standard_sections: list = None,
) -> bytes:
    """Amend the user's own uploaded document: fill-in placeholders for
    missing fields are appended at the end, while missing sections are
    inserted as highlighted headings at their correct position.

    Each missing section is anchored to the nearest FOLLOWING present
    section and inserted right BEFORE it — never to the nearest preceding
    section (inserting "after" a preceding section's heading lands in the
    middle of that section's own content, e.g. before its data table,
    which is wrong, and there's no reliable way to know where a section's
    content ends in order to insert after all of it). Anchoring forward
    always lands at a true section boundary. Falls back to listing at the
    end of the document if nothing present follows (e.g. the missing
    section is meant to be last) or if section_matches/standard_sections
    aren't supplied.
    """
    doc = _load_docx(contents)
    missing_sections = missing_sections or []
    unplaced_sections = list(missing_sections)

    if section_matches and standard_sections:
        heading_index = _heading_index_from_doc(doc, include_tables=True)
        entry_by_id = {entry["id"]: entry for entry in heading_index}
        unplaced_sections = []

        def anchor_element_and_parent(entry):
            if "paragraph" in entry:
                return entry["paragraph"]._p, entry["paragraph"]._parent
            return entry["table"]._tbl, doc

        for section_name in standard_sections:
            if section_name not in missing_sections:
                continue

            idx = standard_sections.index(section_name)
            heading_text = f'Missing Section — "{section_name}" should go here. Please add this section.'

            forward_entry = None
            for i in range(idx + 1, len(standard_sections)):
                next_id = section_matches.get(standard_sections[i])
                if next_id and next_id in entry_by_id:
                    forward_entry = entry_by_id[next_id]
                    break

            if forward_entry is not None:
                anchor_element, parent = anchor_element_and_parent(forward_entry)
                _insert_highlighted_paragraph_before(anchor_element, parent, heading_text)
                continue

            # No following present section to anchor before — we can't
            # reliably tell where a preceding section's own content ends
            # (that's exactly the mid-content bug forward-anchoring
            # avoids), so a missing trailing section falls back to the
            # end-of-document listing rather than guessing a position.
            unplaced_sections.append(section_name)

    if missing_fields or unplaced_sections:
        doc.add_page_break()
        try:
            doc.add_heading("Missing Information — Please Complete", level=1)
        except KeyError:
            # Document has no "Heading 1" style defined — fall back to bold text.
            heading_paragraph = doc.add_paragraph()
            heading_paragraph.add_run("Missing Information — Please Complete").bold = True
        doc.add_paragraph(
            "The fields below were not found in your submission. "
            "Please replace each highlighted instruction with the correct value."
        )

        for field in missing_fields:
            label = FIELD_LABELS.get(field, field.replace("_", " ").title())
            placeholder = FIELD_PLACEHOLDERS.get(field, f"[{label} — enter value here]")
            paragraph = doc.add_paragraph()
            paragraph.add_run(f"The {label} is missing — please include: ").bold = True
            placeholder_run = paragraph.add_run(placeholder)
            placeholder_run.font.highlight_color = WD_COLOR_INDEX.YELLOW

        for section in unplaced_sections:
            paragraph = doc.add_paragraph()
            paragraph.add_run("Missing section — ").bold = True
            section_run = paragraph.add_run(f'please add a "{section}" section to your AOR.')
            section_run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# =========================================================
# PROMPTS
# =========================================================

# Fallback list, used only if the reference AOR template can't be read or
# parsed — the primary source of truth is get_template_reference_sections().
STANDARD_AOR_SECTIONS = [
    "Purpose",
    "Background",
    "Need For Deployment",
    "Scope of Work",
    "Estimated Costs",
    "Net Economic Value (NEV) Analysis and Manpower Capitalisation",
    "Funding",
    "Approval",
    "Annex A: Detailed Scope of Works and Project Timeline",
    "Annex B: Detailed Cost Breakdown and Assessment",
    "Annex C: Net Economic Value (NEV)",
    "Annex D: Finance Manual Reference",
]

_template_reference_sections_cache = None


def get_template_reference_sections() -> list:
    """Reference section list for Document Structure Review, derived from
    the AOR template (AOR_TEMPLATE_PATH). Falls back to
    STANDARD_AOR_SECTIONS if the template is missing/unreadable or yields
    no usable headings. Computed once per process and cached, since the
    template doesn't change between requests.
    """
    global _template_reference_sections_cache
    if _template_reference_sections_cache is not None:
        return _template_reference_sections_cache

    sections = STANDARD_AOR_SECTIONS
    try:
        contents = AOR_TEMPLATE_PATH.read_bytes()
        candidates = extract_candidate_headings(contents, include_tables=False)
        cleaned = _clean_template_heading_candidates(candidates)
        if cleaned:
            sections = cleaned
    except Exception as e:
        logging.warning("template reference sections unavailable, using fallback: %s", e)

    _template_reference_sections_cache = sections
    return sections


_ANNEX_RE = re.compile(r"\bannex\b", re.IGNORECASE)


def _find_annex_start(heading_index: list):
    """Index into heading_index of the first candidate that looks like an
    Annex heading, or None if the document has no recognizable Annex.
    """
    return next(
        (i for i, entry in enumerate(heading_index) if _ANNEX_RE.search(entry["text"])),
        None
    )


def build_structure_review_prompt(heading_index: list, standard_sections: list) -> str:
    annex_start = _find_annex_start(heading_index)
    heading_lines = []
    for i, entry in enumerate(heading_index):
        if i == annex_start:
            heading_lines.append("--- EVERYTHING BELOW THIS LINE IS INSIDE AN ANNEX ---")
        heading_lines.append(f"- {entry['id']}: {entry['text']}")
    headings_text = "\n".join(heading_lines)
    standard_text = "\n".join(f"- {s}" for s in standard_sections)
    return f"""
You are reviewing whether an ICT AOR (Approval of Requirements) Word document
follows the standard CAAS AOR structure.

Standard AOR sections expected (in this rough order):
{standard_text}

Short/standalone lines found in the uploaded document (candidate section
headings, each prefixed with its id, in document order — not every line
here is actually a heading, use judgement to tell real section headings
apart from other short text):
{headings_text}

Sections whose standard name itself starts with "Annex" may be matched to
content anywhere, including below the "EVERYTHING BELOW THIS LINE IS
INSIDE AN ANNEX" marker. Every OTHER standard section (i.e. any main-body
section, not itself an annex) must be matched to a candidate heading that
appears ABOVE that marker. Content that only appears inside an Annex does
NOT satisfy a main-body section's requirement, even if it discusses the
exact same topic in detail — e.g. if the main body never has its own
"Proposed Budget"/"Estimated Costs" section, but Annex B repeats a
detailed CAPEX/OPEX cost breakdown, "Proposed Budget" must still be
reported as missing, because the main body itself never presented it. The
Annex is supplementary backup detail, not a substitute for the main body
actually having the section.

For each standard section, decide whether it is present: only count it as
present if one of the candidate lines above is itself a heading/title for
that section (matched by meaning, not exact wording — e.g. "10 Scope of
Work" counts as "Scope of Work"). Do NOT count a section as present just
because related content is mentioned in passing elsewhere without its own
heading — a heading that has been removed, leaving only orphaned prose
behind, must still be reported as missing.

Known synonym pairs — treat these as ALWAYS the same section regardless of
which wording the document uses (confirmed real-world renamings, not just
a general similarity heuristic):
- "Funding" = "Availability of Funds"
- "Estimated Costs" = "Proposed Budget" = "Cost Breakdown"
- "Need for <xx>" = "Need For Deployment" = "Need for" (any "Need for ..."
  heading satisfies this section, regardless of what follows "for")

Be lenient about matching a heading to a standard section by CONCEPT, not
just similar wording — AOR documents commonly use different titles for the
same section. This leniency has ONE specific, narrow trigger: a candidate
line that literally includes the annotation "(followed by a table
mentioning: ...)" naming CAPEX/OPEX/cost figures — that annotation means a
structured cost-breakdown TABLE immediately follows that heading, which is
strong evidence the heading itself IS a cost/budget disclosure section
under a different title. For example, a candidate "8 Estimated Costs
(followed by a table mentioning: CAPEX, OPEX, Total CAPEX, Total OPEX)"
DOES satisfy a standard section named "Proposed Budget", even though the
two titles share no words.

Do NOT apply this leniency just because a plain paragraph MENTIONS cost
figures, CAPEX, OPEX, or dollar amounts in passing while discussing a
different topic — that is normal narrative, not a cost/budget section.
For example, a heading like "Net Economic Value (NEV) Analysis" followed
by a sentence that cites CAPEX/OPEX numbers to explain a calculation is
NEV analysis referencing figures from elsewhere, NOT a "Proposed
Budget"/"Estimated Costs" section — do not match it to either. If nothing
in the document has the "(followed by a table mentioning: ...)" signal
for cost/budget, and no heading is directly labelled CAPEX/OPEX/Proposed
Budget/Estimated Costs itself, the cost/budget standard section must be
reported as missing, even if cost figures are mentioned somewhere in
passing.

Each candidate id may be used as the match for AT MOST ONE standard
section — never reuse the same candidate for two different standard
sections.

Also check whether the sections that ARE present appear in a reasonable
order relative to the standard list, or if something is clearly out of
place.

Return ONLY a valid JSON object with this exact structure:
{{
  "section_matches": {{"<standard section name>": "<id of the matching candidate line, or null if missing>", ...}},
  "order_note": "<one sentence if something is out of order, otherwise an empty string>"
}}

`section_matches` MUST have exactly one entry for every standard section
listed above, using the id (e.g. "P12" or "T2R0") exactly as it appears
before the colon in the candidate list — do not invent ids, and do not
include the heading text itself as the value.

Do NOT comment on financial figures or data completeness — that is handled
separately. Focus only on document structure. Do not wrap the JSON in
markdown, and do not include any text outside the JSON object.
"""


def build_extraction_prompt(context: str):
    return f"""
You are analysing an ICT Approval of Requirements (AOR) submission using the CAAS AOR template.

Your task is to extract RAW INPUTS only.
Do NOT perform calculations.
Do NOT infer missing values.
Do NOT invent figures, benefits, assumptions, or dates.
Do NOT substitute a figure from a different field or category just because it seems plausible
(e.g. do not use an OPEX figure for CAPEX, or a total budget figure for either). Each field must
come from text explicitly labelled for that exact field.
Do NOT annualise any hours figures.

Return ONLY a valid JSON object.
Do not wrap the JSON in markdown.
Do not include explanations outside the JSON.

Return exactly this JSON structure:

{{
  "project_title": null,
  "purpose": null,
  "problem_statement": null,
  "capex": null,
  "opex": null,
  "project_duration_years": null,
  "annual_productivity_time_savings_hours": null,
  "num_staff": null,
  "savings_duration_months": null,
  "man_hour_rate": null,
  "annual_manpower_impact_fte": null,
  "annual_benefit": null,
  "capex_items": [],
  "opex_items": [],
  "key_assumptions": [],
  "source_evidence": []
}}

Field guidance:
- "project_title": title or name of the ICT project, as explicitly stated
  (e.g. a title line near the top of the document). Do NOT invent or infer
  a title from a description elsewhere.
- "purpose": what the ICT project seeks to achieve. Use ONLY text that
  appears under a heading/section labelled "Purpose" (or a clear synonym
  like "Objective"). Do NOT use similar-sounding content from a
  DIFFERENT section (e.g. "Background", "Need For Deployment", "Scope of
  Work") just because it describes a related goal — those are distinct
  sections in the AOR template and must not be cross-substituted, even if
  the wording overlaps. If there is no section labelled "Purpose" (or
  synonym) with actual content of its own, use null.
- "problem_statement": the gap, pain point, or problem being addressed.
  Use ONLY text that appears under a heading/section labelled "Background"
  or "Problem Statement" (or a clear synonym). Do NOT use content from a
  different section (e.g. "Purpose", "Need For Deployment") even if it
  touches on a similar theme. If no such section with actual content
  exists, use null.

Worked example (a document where the "Purpose" section was removed, but a
different section, "Need For Deployment", happens to contain similar-
sounding text):

  2   Need For Deployment of the Sample System
      a. Based on the trial results, the system has demonstrated
         tangible efficiency benefits for non-critical administrative
         use. The team has assessed that the capability can be
         implemented to:
      b. assist business units in submitting, approving and tracking
         shared-resource bookings through a common workflow; and
  3   expand reporting capability for planning, trend analysis and
      review of common resource utilisation.

There is no "Purpose" heading anywhere in this example. The CORRECT
extraction is "purpose": null — it would be WRONG to extract "assist
business units in submitting, approving and tracking shared-resource
bookings..." as the purpose just because it sounds like one; that text
belongs to "Need For Deployment", a different section, and does not
satisfy the "purpose" field.
- "capex": one-off capital expenditure (implementation, hardware, cybersecurity, contingency).
  Use a figure only if it is either (a) directly labelled CAPEX or "Capital Expenditure" in the
  same sentence/phrase, or (b) a Sub Total / Grand Total / "Say" figure inside a table section
  that is clearly grouped under a CAPEX / Capital Expenditure header, even if that specific row
  does not repeat the word "CAPEX" itself. Prefer the Grand Total (inclusive of contingency) over
  a Sub Total if both are present. Do NOT use a figure from the OPEX section, a combined/overall
  total covering both CAPEX and OPEX together, a man-hour rate, or any other unrelated figure,
  even if it seems plausible. If a CAPEX section header exists but every cost cell under it is
  blank/empty, that means capex is null — do NOT reach into the OPEX section (or anywhere else)
  for a substitute number just because CAPEX itself has none. If no such CAPEX figure exists,
  use null.
- "opex": total operating expenditure over the project duration (licences, maintenance, cloud support).
  Use a figure only if it is either (a) directly labelled OPEX or "Operating Expenditure" in the
  same sentence/phrase, or (b) a Sub Total / Grand Total / "Say" figure inside a table section
  that is clearly grouped under an OPEX / Operating Expenditure header (this section's line items
  are often described as "operations and maintenance", "recurring", "licences", or "subscription"
  without repeating the word "OPEX" on every row). Prefer the Grand Total (inclusive of
  contingency) over a Sub Total if both are present. Do NOT use a figure from the CAPEX section,
  a combined/overall total covering both CAPEX and OPEX together, a man-hour rate, or any other
  unrelated figure, even if it seems plausible. If an OPEX section header exists but every cost
  cell under it is blank/empty, that means opex is null — do NOT reach into the CAPEX section (or
  anywhere else) for a substitute number just because OPEX itself has none. If no such OPEX
  figure exists, use null.
- "capex_items": an array of the individual CAPEX line items found in a CAPEX-labelled cost
  table/section, each as {{"description": <item description>, "amount": <numeric cost for that
  line item>}}. Only include line items that have their own description (e.g. "Provision of
  cybersecurity services"), NOT Sub Total / Contingency / Grand Total / "Say" summary rows — those
  totals are already captured separately via "capex". Preserve the order the items appear in the
  source. If no itemised CAPEX line items are found, use [].
- "opex_items": same as "capex_items" but for the individual line items found in an OPEX-labelled
  cost table/section. If no itemised OPEX line items are found, use [].

Worked example (a cost table where the CAPEX section header exists but every cost cell under
it is blank, and the OPEX section below it has real figures):

  S/N | Description                                          | Estimated Cost ($)
      | CAPEX                                                |
  1   | Implementation of sample system and associated works |
  2   | Provision of cybersecurity services                  |
      | Sub Total                                            |
      | Contingency (5% of Sub Total)                        |
      | Grand Total                                          |
      | Say                                                  |
      | OPEX                                                 |
  1   | Provision of recurring subscription services         | 522,000
  2   | Operations and maintenance support services           | 884,000
      | Sub Total                                            | 1,406,000
      | Contingency (5% of Sub Total)                        | 70,300
      | Grand Total                                          | 1,476,300
      | Say                                                  | 1.5m

The ONLY correct extraction from this table is "capex": null and "opex": 1476300.
It would be WRONG to set "capex" to 1476300, 1406000, or any other number from the OPEX
section — the CAPEX section has no figures of its own, so capex must be null, even though
that means leaving it blank rather than filling in a plausible-looking nearby number.
- "project_duration_years": project duration in years as an integer.   If the document states a date range (e.g. FY25 to FY28), count the  number of full operational years, not the number of financial year 
  labels. For example, FY25 to FY27 = 3 years, FY25 to FY28 = 4 years.  Prefer an explicitly stated duration (e.g. "3-year contract") over  a derived date range if both are present..
- "annual_productivity_time_savings_hours": the RAW total hours figure as stated in the
  document. Do NOT annualise. Do NOT adjust for number of staff or duration. Prefer an explicit
  time-savings/productivity statement (e.g. "the system is expected to save 720 hours annually")
  if one exists. If no such explicit benefit statement exists but the document contains a
  manpower cost-capitalisation table (e.g. "No. of hours", "Manpower Costs Capitalised"
  columns), use that table's hours figure instead — always populate this field from whichever
  hours figure is available rather than leaving it null when any hours figure exists in the
  document.
- "num_staff": number of staff whose time savings are being measured (e.g. 2 project officers).
  Prefer an explicit time-savings context; otherwise fall back to the staff count from a
  manpower cost-capitalisation table, same as "annual_productivity_time_savings_hours" above.
- "savings_duration_months": duration in months over which the raw hours figure was measured
  (e.g. 9 months). Extract as a number. Prefer an explicit time-savings context; otherwise fall
  back to the duration stated in a manpower cost-capitalisation table.
- "man_hour_rate": the man-hour rate in dollars per hour as explicitly stated in the document
  (e.g. 98). Do NOT default to any standard rate if not found. Prefer a rate explicitly tied to a
  benefit/time-savings context; otherwise fall back to the rate stated in a manpower
  cost-capitalisation table.
- "annual_manpower_impact_fte": annual manpower impact in FTE, only if explicitly stated as a
  final FTE figure in the document.
- "annual_benefit": annual benefit in dollars, only if explicitly stated as a final dollar
  benefit figure in the document.
- "key_assumptions": assumptions explicitly stated in the document.
- "source_evidence": short evidence snippets from the context.

Rules:
- Use null for missing single-value fields.
- Use [] for missing list fields.
- Preserve figures and wording as far as possible.
- Do NOT calculate total benefits, total costs, NPV, BCR, annualised hours, or FTE yourself.

Context:
{context}
"""


def build_explanation_prompt(context: str, inputs: dict, metrics: dict, missing_fields: list):
    return f"""
You are reviewing an ICT Approval of Requirements (AOR) submission using the CAAS AOR template.

Use the extracted inputs and computed metrics below.
Do NOT recompute any numbers.
Do NOT invent missing information.
If a metric is null, explain that it could not be computed because the required inputs were not found.

Extracted inputs:
{json.dumps(inputs, indent=2, ensure_ascii=False)}

Computed metrics:
{json.dumps(metrics, indent=2, ensure_ascii=False)}

Missing fields:
{json.dumps(missing_fields, indent=2, ensure_ascii=False)}

Context:
{context}

Return the assessment in this structure:

1. Purpose and Problem Statement

2. Cost Inputs
- CAPEX
- OPEX (total over project duration)
- Project duration

3. Benefits
- Annual productivity time savings (hours)
- Annual manpower impact (FTE)
- Annual benefit ($)

4. Value Assessment
- Total benefits (PV_benefit)
- Total costs (PV_cost)
- Net present value
- Benefit-cost ratio
- Feedback on completeness of user's submission

Stop after section 4. Do NOT write a "Missing Information" section yourself —
it will be added separately from the authoritative missing fields list above.
"""

def build_aor_drafting_prompt(
    context: str,
    inputs: dict,
    metrics: dict,
    approval_info
):
    return f"""
You are drafting a complete CAAS ICT AOR paper.

Use the extracted inputs below.

Extracted Inputs:
{json.dumps(inputs, indent=2)}

Computed Metrics:
{json.dumps(metrics, indent=2)}

Approval Information:
{approval_info}

Source Context:
{context}

Draft the AOR using this structure:

TITLE

Purpose

Background

Need For Deployment

Scope of Work

Estimated Costs

Net Economic Value (NEV) Analysis and Manpower Capitalisation

Funding

Approving Authority

Approval

Rules:
- Write in formal CAAS paper style.
- Use the extracted figures.
- Use approval_info for the Approving Authority section.
- Do not create Annexes.
- Do not invent figures.
- Return ONLY the completed AOR text.
"""
def build_aor_drafting_prompt(
    context: str,
    inputs: dict,
    metrics: dict,
    approval_info
):
    return f"""
    
You are drafting a complete ICT AOR paper.

Use the extracted inputs below.

Extracted Inputs:
{json.dumps(inputs, indent=2, ensure_ascii=False)}

Computed Metrics:
{json.dumps(metrics, indent=2, ensure_ascii=False)}

Approval Information:
{approval_info}

Context:
{context}

Draft the AOR using this structure:

TITLE

Purpose

Background

Need For Deployment

Scope of Work

Estimated Costs

Net Economic Value (NEV) Analysis and Manpower Capitalisation

Funding

Approving Authority

Approval

Rules:
- Use formal CAAS paper writing style.
- Use figures from the extracted inputs.
- Use approval_info for the Approving Authority section.
- Do not create Annexes.
- Do not invent figures.
- Return only the AOR text.
- In the Net Economic Value section, write only a short assessment sentence stating whether the
  project has a positive or negative Net Economic Value (NEV), based on net_present_value (a
  positive net_present_value means a positive NEV). Do NOT list "Present Value of Benefits:",
  "Present Value of Costs:", "Net Present Value:", "Benefit-Cost Ratio:", "Annual Manpower
  Impact:", or "Annual Benefit:" as text — those computed values are already rendered in the
  accompanying NEV table, so repeating them in the narrative is redundant. If net_present_value
  is null, state that the NEV assessment could not be made because the required information was
  not found.
- In the Estimated Costs section, write only a short introductory sentence stating that the
  estimated costs for the project are set out in the table below (e.g. "The estimated costs for
  the project are as follows:"). Do NOT list "CAPEX:"/"OPEX:" labels, individual line items, Sub
  Total, Contingency, or Grand Total figures as text — that breakdown is already rendered in the
  accompanying cost breakdown table, so repeating it in the narrative is redundant.

"""

# =========================================================
# MAIN ENDPOINT
# =========================================================

@app.post("/process")
async def process(query: str = Form(""), file: UploadFile = File(None)):

    if file:
        if not file.filename.lower().endswith(".docx"):
            raise HTTPException(
                status_code=400,
                detail="Please upload your AOR as a Word document (.docx) file."
            )

        contents = await file.read()

        raw_docs = extract_documents(contents, file.filename)

        if not raw_docs:
            raise HTTPException(
                status_code=400,
                detail="The document appears to be empty or unreadable."
            )
   
    else:
        if not query.strip():
            raise HTTPException(
                 status_code=400,
                 detail="Please provide either a file or some text."
            )

        raw_docs = [
            LCDocument(
                page_content=query,
                metadata={
                    "filename":"user_prompt",
                    "source_type":"prompt",
                    "content_type":"plain_text"
                }
            )
        ]

    embeddings = OpenAIEmbeddings()
    docs = split_documents(raw_docs, embeddings)

    if not docs:
        raise HTTPException(
            status_code=400,
            detail="No readable content could be extracted from the document."
        )

    store = FAISS.from_documents(docs, embeddings)
    retriever = store.as_retriever(search_kwargs={"k": 8})

    retrieval_query = query.strip() or (
        "project objective business need benefits "
        "cost estimate capex opex manpower impact "
        "fte time saving annual benefit pv factor "
        "implementation period project duration"
    )

    relevant_docs = retriever.invoke(retrieval_query)

    context_blocks = []
    for i, doc in enumerate(relevant_docs, start=1):
        metadata = doc.metadata
        source_parts = []
        if metadata.get("filename"):
            source_parts.append(f"File: {metadata.get('filename')}")
        if metadata.get("page"):
            source_parts.append(f"Page: {metadata.get('page')}")
        if metadata.get("sheet"):
            source_parts.append(f"Sheet: {metadata.get('sheet')}")
        if metadata.get("table_index"):
            source_parts.append(f"Table: {metadata.get('table_index')}")
        if metadata.get("row_index"):
            source_parts.append(f"Row: {metadata.get('row_index')}")
        if metadata.get("chunk_index"):
            source_parts.append(f"Chunk: {metadata.get('chunk_index')}")
        source_label = " | ".join(source_parts)
        context_blocks.append(f"[Chunk {i} | {source_label}]\n{doc.page_content}")

    context = "\n\n".join(context_blocks)

    llm = ChatOpenAI(model="gpt-4o-mini", temperature=0)

    extraction_prompt = build_extraction_prompt(context)
    extraction_response = llm.invoke([
        SystemMessage(content=extraction_prompt),
        HumanMessage(content="Extract ICT AOR inputs from the context.")
    ])

    try:
        extracted_inputs_raw = parse_json_from_llm(extraction_response.content)
        extracted_inputs = clean_extracted_inputs(extracted_inputs_raw)
    except Exception as e:
        return {
            "status": "error",
            "message": "Failed to extract structured ICT AOR inputs.",
            "details": str(e),
            "raw_llm_response": extraction_response.content
        }

    computed_metrics = compute_aor_metrics(extracted_inputs)
    approval_info = determine_approving_authority(
        compute_cost_totals(extracted_inputs, computed_metrics)["grand_total"]
    )
    missing_fields = identify_missing_fields(extracted_inputs, computed_metrics)
    submission_complete = len(missing_fields) == 0

    aor_prompt = build_aor_drafting_prompt(
        context=context,
        inputs=extracted_inputs,
        metrics=computed_metrics,
        approval_info=approval_info
    )

    final_response = llm.invoke([
        SystemMessage(content=aor_prompt),
        HumanMessage(content="Draft the AOR.")
    ])

    # docx_source_text stays intro-only for "Estimated Costs"/"NEV" — the
    # docx gets its figures from the deterministic tables instead (built
    # from extracted_inputs/computed_metrics, see build_result_docx). The
    # bullet breakdown below is only for the chat/preview response; it must
    # NOT also flow into the docx, since its "- label: value" line format
    # doesn't match strip_cost_breakdown_narrative/strip_nev_metrics_narrative
    # (which target the LLM's own numbered/"Sub Total:"-style phrasing) and
    # would otherwise leak into the "Proposed Budget" paragraph as duplicate
    # text above the real table.
    docx_source_text = strip_nev_metrics_narrative(
        strip_cost_breakdown_narrative(final_response.content.rstrip())
    )
    full_result = insert_deterministic_breakdown(
        docx_source_text,
        ["Proposed Budget", "Estimated Costs"],
        build_cost_breakdown_bullets(extracted_inputs, computed_metrics)
    )
    full_result = insert_deterministic_breakdown(
        full_result,
        [
            "Net Economic Value (NEV) Analysis and Manpower Capitalisation",
            "Net Economic Value Analysis and Manpower Capitalisation"
        ],
        build_nev_bullets(computed_metrics)
    )
    if missing_fields:
        full_result += "\n\n" + build_missing_information_section(missing_fields)

    missing_sections = []
    section_matches = {}
    standard_sections = []
    if file and file.filename.lower().endswith(".docx"):
        try:
            standard_sections = get_template_reference_sections()
            heading_index = annotate_heading_index_synonyms(
                build_heading_index(contents, with_context=True)
            )
            structure_prompt = build_structure_review_prompt(heading_index, standard_sections)
            structure_response = llm.invoke([
                SystemMessage(content=structure_prompt),
                HumanMessage(content="Assess the document's structure against the standard AOR format.")
            ])
            structure_data = parse_json_from_llm(structure_response.content)
            if isinstance(structure_data, dict):
                section_matches = structure_data.get("section_matches") or {}
                section_matches = resolve_cost_section_matches(section_matches, heading_index)
                section_matches = deduplicate_reused_matches(section_matches, heading_index)
                order_note = structure_data.get("order_note") or ""
            else:
                order_note = ""
            missing_sections = [
                section for section in standard_sections
                if not section_matches.get(section)
            ]
            full_result += "\n\n" + build_structure_review_section(missing_sections, order_note)
        except Exception:
            pass

    amended_docx_base64 = None
    amended_docx_filename = None

    import logging
    logging.warning(f"EXTRACTED INPUTS: {json.dumps(extracted_inputs, indent=2)}")
    logging.warning(f"COMPUTED METRICS: {json.dumps(computed_metrics, indent=2)}")

    if file and (missing_fields or missing_sections):
        amended_docx_base64 = base64.b64encode(
            build_amended_docx(
                contents, missing_fields, missing_sections,
                section_matches=section_matches, standard_sections=standard_sections
            )
        ).decode("ascii")
        amended_docx_filename = f"amended-{file.filename}"
    elif not file:
        amended_docx_base64 = base64.b64encode(
            build_result_docx(docx_source_text, inputs=extracted_inputs, metrics=computed_metrics)
        ).decode("ascii")
        amended_docx_filename = "aor-assessment.docx"

    return {
        "status": "ok",
        "approval_info": approval_info,
        "submission_complete": submission_complete,
        "result": full_result,
        "extracted_inputs": extracted_inputs,
        "computed_metrics": computed_metrics,
        "missing_fields": missing_fields,
        "amended_docx_base64": amended_docx_base64,
        "amended_docx_filename": amended_docx_filename,
        "sources_used": [
            {
                "metadata": doc.metadata,
                "preview": doc.page_content[:300]
            }
            for doc in relevant_docs
        ]
    }